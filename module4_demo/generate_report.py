"""
generate_report.py - Synthetic Sentinel Final Report Generator
==============================================================
Phase 5: Programmatically generate the final report as a Word document.

Usage:
    python generate_report.py
"""

import json
import os
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(PROJECT_DIR, "data")
RESULTS_DIR = os.path.join(PROJECT_DIR, "results")
MODEL_DIR = os.path.join(PROJECT_DIR, "models")


def load_json_if_exists(path):
    """Load JSON when available, otherwise return an empty dict."""
    if not os.path.exists(path):
        return {}
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def set_paragraph_format(paragraph, space_after=6, space_before=0, line_spacing=1.15):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing


def add_heading_styled(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
    return heading


def add_image_if_exists(doc, path, width=Inches(5.5), caption=None):
    """Add an image with an optional caption."""
    if not os.path.exists(path):
        return False

    doc.add_picture(path, width=width)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True
        cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return True


def add_metric_table(doc, baseline_metrics, hybrid_metrics):
    """Render the model metrics table."""
    table = doc.add_table(rows=3, cols=6)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Model", "Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    table.rows[1].cells[0].text = "Logistic Regression"
    for i, key in enumerate(["accuracy", "precision", "recall", "f1", "roc_auc"], 1):
        table.rows[1].cells[i].text = f"{baseline_metrics.get(key, 0):.4f}"

    table.rows[2].cells[0].text = "Hybrid (RoBERTa)"
    for i, key in enumerate(["accuracy", "precision", "recall", "f1", "roc_auc"], 1):
        table.rows[2].cells[i].text = f"{hybrid_metrics.get(key, 0):.4f}"


def describe_baseline_comparison(results):
    """Create a comparison paragraph that matches the actual evaluation artifacts."""
    baseline = results.get("baseline_metrics", {})
    hybrid = results.get("hybrid_metrics", {})
    baseline_f1 = baseline.get("f1")
    hybrid_f1 = hybrid.get("f1")

    if baseline_f1 is None or hybrid_f1 is None:
        return "Baseline and hybrid metrics were not both available when this report was generated."

    if hybrid_f1 > baseline_f1:
        return (
            f"In the current evaluation artifacts, the hybrid model outperforms the baseline on F1 "
            f"({hybrid_f1:.4f} vs {baseline_f1:.4f})."
        )

    if hybrid_f1 < baseline_f1:
        return (
            f"In the current evaluation artifacts, the baseline outperforms the hybrid model on F1 "
            f"({baseline_f1:.4f} vs {hybrid_f1:.4f}). This usually means the current dataset is "
            "dominated by easy lexical or template cues, so the benchmark is not yet stressing the "
            "model in a realistic way."
        )

    return f"In the current evaluation artifacts, the baseline and hybrid model tie on F1 ({hybrid_f1:.4f})."


def create_report():
    """Generate the full final report."""
    results = load_json_if_exists(os.path.join(RESULTS_DIR, "evaluation_results.json"))
    dataset_metadata = load_json_if_exists(os.path.join(DATA_DIR, "dataset_metadata.json"))
    training_results = load_json_if_exists(os.path.join(MODEL_DIR, "training_results.json"))

    doc = Document()

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Synthetic Sentinel")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("A Multi-Model Approach to AI-Generated Misinformation Detection")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(0x58, 0x58, 0x58)

    doc.add_paragraph()

    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team_run = team.add_run("Group 9")
    team_run.font.size = Pt(14)
    team_run.font.bold = True

    members = [
        "Angel Daniel Bustamante Perez - Data Architect",
        "Romilson Lemes Cordeiro - ML Engineer",
        "Sakthivel Rithiek - Linguistic Analyst",
        "Jack Si - UX and Business Lead",
    ]
    for member in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(member).font.size = Pt(11)

    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course_run = course.add_run(
        f"ARTI407 - NLP Final Project\nWinter 2026\n{datetime.now().strftime('%B %d, %Y')}"
    )
    course_run.font.size = Pt(11)
    course_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Planning and Proposal",
        "2. Implementation",
        "3. Analysis and Results",
        "4. Conclusion and Future Work",
        "5. References",
    ]
    for item in toc_items:
        paragraph = doc.add_paragraph(item)
        set_paragraph_format(paragraph, space_after=2)

    doc.add_page_break()

    add_heading_styled(doc, "1. Planning and Proposal", level=1)
    add_heading_styled(doc, "1.1 Project Goal and Problem Statement", level=2)
    doc.add_paragraph(
        "The goal of this project is to build an NLP pipeline that distinguishes human-written text "
        "from AI-generated content. The project combines a transformer encoder with lightweight "
        "stylometric features to produce a probability of synthetic authorship for free-form text."
    )

    add_heading_styled(doc, "1.2 Dataset Description", level=2)
    doc.add_paragraph(
        "The intended dataset design combines social media, news-like text, and custom synthetic "
        "samples so the detector can learn across multiple writing styles."
    )

    table = doc.add_table(rows=4, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Source", "Nominal Size", "Type", "Description"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    rows = [
        ["TweepFake", "25,000+", "Tweets", "Human vs bot-generated tweets from HuggingFace"],
        ["FakeNewsNet", "3,000", "News Articles", "Local CSV if provided, synthetic fallback otherwise"],
        ["Custom Synthetic", "1,000", "Mixed", "Locally generated human-like and AI-like samples"],
    ]
    for r, row in enumerate(rows, 1):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = value

    if dataset_metadata:
        synthetic_fraction = dataset_metadata.get("synthetic_fraction", 0.0) * 100
        doc.add_paragraph()
        doc.add_paragraph(
            "Current repository build note: the generated splits in this workspace rely heavily on "
            f"synthetic fallback data ({synthetic_fraction:.1f}% of samples). The resulting metrics "
            "should be read as synthetic/template benchmark performance rather than broad real-world "
            "generalization."
        )

    add_heading_styled(doc, "1.3 Evaluation Metrics", level=2)
    metrics = [
        "F1-Score as the primary balance between precision and recall",
        "Accuracy for overall correctness",
        "Precision and recall to track false positives and false negatives separately",
        "ROC-AUC for threshold-independent separability",
    ]
    for item in metrics:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    add_heading_styled(doc, "2. Implementation", level=1)
    add_heading_styled(doc, "2.1 Technical Architecture", level=2)
    doc.add_paragraph(
        "The project is organized into modular scripts for data processing, training, evaluation, "
        "demo delivery, and report generation."
    )
    modules = [
        ("data_loader.py", "Builds train/validation/test CSVs and feature columns."),
        ("train_model.py", "Trains the TF-IDF baseline and the RoBERTa hybrid classifier."),
        ("evaluate.py", "Evaluates saved models and writes metrics plus analysis plots."),
        ("app.py", "Provides a Streamlit demo for single-text and batch inference."),
        ("generate_report.py", "Builds this Word report from project artifacts."),
    ]
    for name, description in modules:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(description)

    add_heading_styled(doc, "2.2 Team Module Ownership", level=2)
    ownership = [
        "Module 1 - Data Pipeline and Dataset Governance: Data Architect",
        "Module 2 - Modeling and Training: ML Engineer",
        "Module 3 - Evaluation and Linguistic Analysis: Linguistic Analyst",
        "Module 4 - Demo and Reporting: UX and Business Lead",
    ]
    for item in ownership:
        doc.add_paragraph(item, style="List Bullet")

    add_heading_styled(doc, "2.3 Data Pipeline", level=2)
    pipeline_steps = [
        "Dataset acquisition with synthetic fallback when remote or local sources are unavailable",
        "Text cleaning and exact-duplicate removal",
        "Burstiness feature computation from sentence-length variance",
        "GPT-2 perplexity feature computation",
        "Stratified 70/15/15 train/validation/test split",
        "Dataset metadata export for source-composition auditing",
    ]
    for step in pipeline_steps:
        doc.add_paragraph(step, style="List Bullet")

    add_heading_styled(doc, "2.4 Model Training", level=2)
    doc.add_paragraph(
        "Two models are trained for comparison: a Logistic Regression baseline over TF-IDF plus "
        "statistical features, and a HybridDetector that concatenates RoBERTa's [CLS] embedding "
        "with perplexity and burstiness before classification."
    )
    doc.add_paragraph(
        "The training script now respects command-line hyperparameters, supports optional encoder "
        "freezing for fast experiments, and uses a linear warmup scheduler instead of documenting a "
        "scheduler that was never applied."
    )
    if training_results.get("hybrid"):
        doc.add_paragraph(
            f"Current run summary: `training_results.json` reports a best validation F1 of "
            f"{training_results['hybrid'].get('best_f1', 0):.4f} for the hybrid model."
        )

    doc.add_page_break()

    add_heading_styled(doc, "3. Analysis and Results", level=1)
    add_heading_styled(doc, "3.1 Model Performance", level=2)
    if results:
        add_metric_table(
            doc,
            results.get("baseline_metrics", {}),
            results.get("hybrid_metrics", {}),
        )
    else:
        doc.add_paragraph("[Run `python evaluate.py` to populate the metrics table.]")

    add_heading_styled(doc, "3.2 Confusion Matrix", level=2)
    add_image_if_exists(
        doc,
        os.path.join(RESULTS_DIR, "confusion_matrix.png"),
        caption="Figure 1: Confusion Matrix - Hybrid Detector",
    )

    add_heading_styled(doc, "3.3 Probability Distribution", level=2)
    doc.add_paragraph(
        "This plot helps visualize how confidently the hybrid model separates the two classes on the "
        "current benchmark. Strong separation on a mostly synthetic dataset should be interpreted "
        "cautiously."
    )
    add_image_if_exists(
        doc,
        os.path.join(RESULTS_DIR, "prob_dist.png"),
        caption="Figure 2: Prediction Probability Distribution",
    )

    add_heading_styled(doc, "3.4 Burstiness Analysis", level=2)
    doc.add_paragraph(
        "Burstiness remains a useful descriptive feature, but on template-heavy data it can become a "
        "proxy for template style rather than authentic human variability."
    )
    add_image_if_exists(
        doc,
        os.path.join(RESULTS_DIR, "burstiness.png"),
        caption="Figure 3: Burstiness vs Prediction Probability",
    )

    add_heading_styled(doc, "3.5 Baseline Comparison", level=2)
    if results:
        doc.add_paragraph(describe_baseline_comparison(results))
    else:
        doc.add_paragraph("Baseline comparison will be available after running `evaluate.py`.")
    add_image_if_exists(
        doc,
        os.path.join(RESULTS_DIR, "training_curves.png"),
        caption="Figure 4: Training and Validation Curves",
    )
    add_image_if_exists(
        doc,
        os.path.join(RESULTS_DIR, "roc_curve.png"),
        caption="Figure 5: ROC Curve - Hybrid Detector",
    )

    add_heading_styled(doc, "3.6 Error Analysis and Limitations", level=2)
    doc.add_paragraph(
        "The most important limitation in the current repository state is dataset realism. Because the "
        "generated splits rely heavily on synthetic fallback data, high scores may reflect template "
        "memorization or lexical shortcut learning instead of robust authorship detection."
    )
    limitations = [
        "Most currently generated split files are synthetic fallback data",
        "Short texts provide limited signal for reliable classification",
        "Human-edited AI text remains a gray zone for both lexical and contextual detectors",
        "Perplexity is computed with GPT-2, which is a weak proxy for newer language models",
        "A baseline beating the hybrid model is a sign that the benchmark is too easy or too templated",
    ]
    for item in limitations:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    add_heading_styled(doc, "4. Conclusion and Future Work", level=1)
    doc.add_paragraph(
        "Synthetic Sentinel is now a more honest and internally consistent course project: the demo "
        "fails safely without weights, the training script matches its documented controls, and the "
        "report reflects the actual artifacts in the repository. The next major step is to rebuild the "
        "benchmark around more real-source data so the reported metrics become meaningful beyond the "
        "synthetic setting."
    )
    future_work = [
        "Restore real-source datasets and rerun the full pipeline",
        "Add source-stratified or source-held-out evaluation",
        "Test on human-edited AI text rather than only clean class labels",
        "Replace or augment GPT-2 perplexity with stronger detection features",
        "Package inference as a service once the benchmark is more realistic",
    ]
    for item in future_work:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    add_heading_styled(doc, "5. References", level=1)
    references = [
        "Liu, Y., et al. (2019). RoBERTa: A Robustly Optimized BERT Pretraining Approach.",
        "Fagni, T., et al. (2021). TweepFake: About Detecting Deepfake Tweets.",
        "Shu, K., et al. (2020). FakeNewsNet: A Data Repository with News Content, Social Context, and Spatiotemporal Information.",
        "Radford, A., et al. (2019). Language Models are Unsupervised Multitask Learners.",
        "Mitchell, E., et al. (2023). DetectGPT: Zero-Shot Machine-Generated Text Detection using Probability Curvature.",
        "Gehrmann, S., et al. (2019). GLTR: Statistical Detection and Visualization of Generated Text.",
    ]
    for i, ref in enumerate(references, 1):
        paragraph = doc.add_paragraph(f"[{i}] {ref}")
        set_paragraph_format(paragraph, space_after=4)

    output_path = os.path.join(PROJECT_DIR, "Final_Report.docx")
    doc.save(output_path)
    print(f"\nFinal report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_report()
